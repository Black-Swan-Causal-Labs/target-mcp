"""Server-level integration tests: the MCP tool surface, structured I/O, and
the interactive review round-trip through call_tool()."""

import asyncio

import pytest

from target_mcp.server import mcp

FAKE = """Emulating a target trial of drug A vs B
Abstract
We emulated a target trial using claims data. Hazard ratio 0.85 under no unmeasured confounding.
Methods
Eligibility required age over 40. Follow-up started at treatment assignment (time zero) and ended at death.
We adjusted for baseline confounders using inverse probability weighting.
Results
Of 120000 assessed, 45000 eligible. The hazard ratio was 0.85.
Discussion
Residual confounding remains possible.
"""

EXPECTED_TOOLS = {
    "get_checklist", "parse_manuscript", "parse_pmcid", "assess_manuscript",
    "submit_scaffold_verdicts", "render_checklist", "render_checklist_docx",
    "render_checklist_html", "aggregate_corpus", "build_coding_sheet",
    "validate_against_gold",
}


def _call(name, args):
    # call_tool returns (content_blocks, structured_content)
    return asyncio.run(mcp.call_tool(name, args))[1]


def test_tool_surface_and_output_schemas():
    tools = asyncio.run(mcp.list_tools())
    assert {t.name for t in tools} == EXPECTED_TOOLS
    # every tool declares a structured output schema
    for t in tools:
        assert getattr(t, "outputSchema", None) is not None, t.name
    # submit_scaffold_verdicts takes a structured array, not a JSON string
    sub = next(t for t in tools if t.name == "submit_scaffold_verdicts")
    assert sub.inputSchema["properties"]["items"]["type"] == "array"


def test_review_round_trip_structured_io():
    parsed = _call("parse_manuscript", {"document": FAKE, "manuscript_id": "rt-1"})
    sha = parsed["text_sha256"]
    assert parsed["n_pages"] is None  # text input -> nullable field tolerated
    assert {"methods", "results"} <= {s["name"] for s in parsed["sections"]}

    scaf = _call("assess_manuscript", {"document": sha, "mode": "scaffold"})
    assert scaf["mode"] == "scaffold" and len(scaf["leaf_ids"]) == 39

    items = [
        {"id": lid,
         "verdict": "reported" if lid == "6d" else "not_reported",
         "confidence": 0.8, "rationale": "x",
         **({"evidence_quotes": ["Follow-up started at treatment assignment (time zero)"]}
            if lid == "6d" else {})}
        for lid in scaf["leaf_ids"]
    ]
    res = _call("submit_scaffold_verdicts", {"text_sha256": sha, "items": items, "model": "test"})
    assert res["mode"] == "scaffold" and len(res["items"]) == 39
    r6d = next(i for i in res["items"] if i["id"] == "6d")
    assert r6d["evidence"][0]["resolved"] and r6d["evidence"][0]["source_document"] == "main"

    agg = _call("aggregate_corpus", {"use_session": True})
    assert agg["n_papers"] >= 1


def test_render_checklist_round_trip():
    parsed = _call("parse_manuscript", {"document": FAKE, "manuscript_id": "rt-3",
                                        "supplement_status": "none_exists"})
    sha = parsed["text_sha256"]
    scaf = _call("assess_manuscript", {"document": sha, "mode": "scaffold"})
    items = [
        {"id": lid,
         "verdict": "reported" if lid == "6d" else "not_reported",
         "confidence": 0.8, "rationale": "x",
         **({"evidence_quotes": ["Follow-up started at treatment assignment (time zero)"]}
            if lid == "6d" else {})}
        for lid in scaf["leaf_ids"]
    ]
    _call("submit_scaffold_verdicts", {"text_sha256": sha, "items": items, "model": "test"})

    # render by manuscript_id from the session
    report = _call("render_checklist", {"manuscript_id": "rt-3"})
    assert len(report["rows"]) == 39
    assert report["view"] == "enriched"
    assert "floor" not in report
    assert "(TARGET) Checklist" in report["markdown"]
    row6d = next(r for r in report["rows"] if r["id"] == "6d")
    assert row6d["location"] == "Methods"


def test_render_checklist_docx_round_trip(tmp_path):
    parsed = _call("parse_manuscript", {"document": FAKE, "manuscript_id": "rt-4",
                                        "supplement_status": "none_exists"})
    sha = parsed["text_sha256"]
    scaf = _call("assess_manuscript", {"document": sha, "mode": "scaffold"})
    items = [
        {"id": lid,
         "verdict": "reported" if lid == "6d" else "not_reported",
         "confidence": 0.8, "rationale": "x",
         **({"evidence_quotes": ["Follow-up started at treatment assignment (time zero)"]}
            if lid == "6d" else {})}
        for lid in scaf["leaf_ids"]
    ]
    _call("submit_scaffold_verdicts", {"text_sha256": sha, "items": items, "model": "test"})

    out = str(tmp_path / "rt4.docx")
    res = _call("render_checklist_docx", {"manuscript_id": "rt-4", "path": out})
    assert res["path"] == out and res["view"] == "enriched" and res["n_rows"] == 39
    assert res["bytes"] > 0

    # the file is a valid docx whose one table carries every leaf id
    from docx import Document
    doc = Document(out)
    seen = {c.text.strip()
            for t in doc.tables for row in t.rows for c in [row.cells[0]]}
    assert {"6d", "7g.i", "1a", "21"} <= seen


def test_render_docx_logo_gated_to_official_view(tmp_path):
    parsed = _call("parse_manuscript", {"document": FAKE, "manuscript_id": "rt-5",
                                        "supplement_status": "none_exists"})
    sha = parsed["text_sha256"]
    scaf = _call("assess_manuscript", {"document": sha, "mode": "scaffold"})
    items = [{"id": lid, "verdict": "not_reported", "confidence": 0.7, "rationale": "x"}
             for lid in scaf["leaf_ids"]]
    _call("submit_scaffold_verdicts", {"text_sha256": sha, "items": items, "model": "t"})

    # logo requested on the enriched (third-party) view -> ignored
    enr = _call("render_checklist_docx", {"manuscript_id": "rt-5", "view": "enriched",
                                          "logo": True, "path": str(tmp_path / "e.docx")})
    assert enr["logo_applied"] is False

    # logo on the official (self-report) view -> applied, image embedded
    off = _call("render_checklist_docx", {"manuscript_id": "rt-5", "view": "official",
                                          "logo": True, "path": str(tmp_path / "o.docx")})
    assert off["logo_applied"] is True
    import zipfile
    with zipfile.ZipFile(off["path"]) as z:
        assert any(n.startswith("word/media/") for n in z.namelist())
    # off by default even on the official view
    plain = _call("render_checklist_docx", {"manuscript_id": "rt-5", "view": "official",
                                            "path": str(tmp_path / "p.docx")})
    assert plain["logo_applied"] is False


def test_render_docx_default_path_is_writable():
    import os
    import tempfile
    parsed = _call("parse_manuscript", {"document": FAKE, "manuscript_id": "wr"})
    sha = parsed["text_sha256"]
    scaf = _call("assess_manuscript", {"document": sha, "mode": "scaffold"})
    items = [{"id": lid, "verdict": "not_reported", "confidence": 0.7, "rationale": "x"}
             for lid in scaf["leaf_ids"]]
    _call("submit_scaffold_verdicts", {"text_sha256": sha, "items": items})
    # no path given -> must default into a writable temp dir, not the filesystem root
    res = _call("render_checklist_docx", {"manuscript_id": "wr"})
    assert res["path"].startswith(tempfile.gettempdir())
    assert os.path.exists(res["path"])


def test_submit_reparses_on_cache_miss():
    from target_mcp.server import _parsed
    parsed = _call("parse_manuscript", {"document": FAKE, "manuscript_id": "cm"})
    sha = parsed["text_sha256"]
    scaf = _call("assess_manuscript", {"document": sha, "mode": "scaffold"})
    items = [{"id": lid, "verdict": "not_reported", "confidence": 0.7, "rationale": "x"}
             for lid in scaf["leaf_ids"]]
    _parsed.clear()  # simulate a server restart wiping the in-process parse cache
    with pytest.raises(Exception):
        _call("submit_scaffold_verdicts", {"text_sha256": sha, "items": items})
    # re-supplying the document re-parses transparently so staged verdicts survive
    res = _call("submit_scaffold_verdicts",
                {"text_sha256": sha, "items": items, "document": FAKE})
    assert len(res["items"]) == 39


def test_submit_returns_stamped_deliverable_inline():
    # The structural safeguard: the finished, stamped render falls out of submit,
    # so there is no separate render step to skip or hand-build around.
    import base64
    parsed = _call("parse_manuscript", {"document": FAKE, "manuscript_id": "deliv",
                                        "supplement_status": "none_exists"})
    sha = parsed["text_sha256"]
    scaf = _call("assess_manuscript", {"document": sha, "mode": "scaffold"})
    items = [{"id": lid, "verdict": "not_reported", "confidence": 0.7, "rationale": "x"}
             for lid in scaf["leaf_ids"]]
    res = _call("submit_scaffold_verdicts", {"text_sha256": sha, "items": items, "model": "t"})
    rep = res["report"]
    assert rep["provenance"]["stamp"].startswith("TGT-")
    # Default bundle is HTML ONLY: extra formats transit the caller's context
    # window (the triple-format bundle caused a real client timeout).
    assert rep["html"].startswith("<!doctype html>")
    assert rep["provenance"]["stamp"] in rep["html"]
    assert "markdown" not in rep and "docx_base64" not in rep
    # Opting in via report_formats still yields the full stamped bundle.
    res = _call("submit_scaffold_verdicts",
                {"text_sha256": sha, "items": items, "model": "t",
                 "report_formats": ["html", "markdown", "docx"]})
    rep = res["report"]
    assert "(TARGET) Checklist" in rep["markdown"] and rep["provenance"]["stamp"] in rep["markdown"]
    assert rep["html"].startswith("<!doctype html>")
    # docx bytes are in hand, valid, and carry the stamp in core properties
    raw = base64.b64decode(rep["docx_base64"])
    assert raw[:2] == b"PK"
    from docx import Document
    import io
    doc = Document(io.BytesIO(raw))
    assert doc.core_properties.keywords == rep["provenance"]["stamp"]


def test_render_checklist_html_tool():
    parsed = _call("parse_manuscript", {"document": FAKE, "manuscript_id": "htool",
                                        "supplement_status": "none_exists"})
    sha = parsed["text_sha256"]
    scaf = _call("assess_manuscript", {"document": sha, "mode": "scaffold"})
    items = [{"id": lid, "verdict": "not_reported", "confidence": 0.7, "rationale": "x"}
             for lid in scaf["leaf_ids"]]
    _call("submit_scaffold_verdicts", {"text_sha256": sha, "items": items, "render": False})
    out = _call("render_checklist_html", {"manuscript_id": "htool"})
    assert out["html"].startswith("<!doctype html>") and out["filename"].endswith(".html")
    import base64
    assert base64.b64decode(out["content_base64"]).startswith(b"<!doctype html>")
    assert out["provenance"]["stamp"].startswith("TGT-")


def test_render_docx_returns_retrievable_bytes():
    import base64
    import io
    from docx import Document
    parsed = _call("parse_manuscript", {"document": FAKE, "manuscript_id": "b64"})
    sha = parsed["text_sha256"]
    scaf = _call("assess_manuscript", {"document": sha, "mode": "scaffold"})
    items = [{"id": lid, "verdict": "not_reported", "confidence": 0.7, "rationale": "x"}
             for lid in scaf["leaf_ids"]]
    _call("submit_scaffold_verdicts", {"text_sha256": sha, "items": items})
    res = _call("render_checklist_docx", {"manuscript_id": "b64"})
    # the deliverable is retrievable without touching the server filesystem
    assert res["content_base64"] and res["filename"].endswith(".docx")
    raw = base64.b64decode(res["content_base64"])
    assert raw[:2] == b"PK"  # zip/docx magic
    doc = Document(io.BytesIO(raw))  # reopens as a valid Word document
    assert any("(TARGET) Checklist" in p.text for p in doc.paragraphs)


def test_submit_usage_field_is_nullable_in_schema():
    # regression for the "None is not of type 'object'" crash: scaffold results
    # omit `usage`, whose total=False default is null; the schema must accept null
    import asyncio
    tools = asyncio.run(mcp.list_tools())
    sub = next(t for t in tools if t.name == "submit_scaffold_verdicts")
    usage = sub.outputSchema["properties"]["usage"]
    types = {b.get("type") for b in usage.get("anyOf", [usage])}
    assert "null" in types and "object" in types


def test_render_checklist_accepts_structured_assessment():
    parsed = _call("parse_manuscript", {"document": FAKE, "manuscript_id": "rt-2",
                                        "supplement_status": "none_exists"})
    sha = parsed["text_sha256"]
    scaf = _call("assess_manuscript", {"document": sha, "mode": "scaffold"})
    items = [{"id": lid, "verdict": "not_reported", "confidence": 0.7, "rationale": "x"}
             for lid in scaf["leaf_ids"]]
    assessment = _call("submit_scaffold_verdicts", {"text_sha256": sha, "items": items})
    # pass the assessment object straight in (structured input, not a JSON string)
    report = _call("render_checklist", {"assessment": assessment})
    assert len(report["rows"]) == 39
    assert report["completeness"]["not_reported"] == 39


def test_citation_flows_from_parse_to_every_render():
    import base64
    apa = ("Doe, J., & Roe, R. (2024). A target trial emulation of drug A "
           "versus drug B. Journal of Testing, 12(3), 1-10. "
           "https://doi.org/10.1000/test")
    parsed = _call("parse_manuscript", {"document": FAKE, "manuscript_id": "cite-1",
                                        "supplement_status": "none_exists",
                                        "citation": apa})
    assert parsed["citation"] == apa
    sha = parsed["text_sha256"]
    scaf = _call("assess_manuscript", {"document": sha, "mode": "scaffold"})
    items = [{"id": lid, "verdict": "not_reported", "confidence": 0.7, "rationale": "x"}
             for lid in scaf["leaf_ids"]]
    res = _call("submit_scaffold_verdicts",
                {"text_sha256": sha, "items": items,
                 "report_formats": ["html", "markdown", "docx"]})
    assert res["citation"] == apa
    rep = res["report"]
    import html as _html
    assert _html.escape(apa) in rep["html"] and "Manuscript assessed:" in rep["html"]
    assert f"**Manuscript assessed:** {apa}" in rep["markdown"]
    from docx import Document
    import io
    doc = Document(io.BytesIO(base64.b64decode(rep["docx_base64"])))
    assert any(apa in p.text for p in doc.paragraphs)
    # without a citation, renders fall back to the short manuscript_id
    parsed2 = _call("parse_manuscript", {"document": FAKE + " x", "manuscript_id":
                                         "cite-2", "supplement_status": "none_exists"})
    scaf2 = _call("assess_manuscript", {"document": parsed2["text_sha256"],
                                        "mode": "scaffold"})
    items2 = [{"id": lid, "verdict": "not_reported", "confidence": 0.7, "rationale": "x"}
              for lid in scaf2["leaf_ids"]]
    res2 = _call("submit_scaffold_verdicts",
                 {"text_sha256": parsed2["text_sha256"], "items": items2})
    assert "Manuscript assessed:</span> cite-2" in res2["report"]["html"]
