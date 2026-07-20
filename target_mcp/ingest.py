"""Ingestion layer: document -> SectionMap with character-offset spans.

The SectionMap is the deterministic substrate every assessment runs over.
Evidence spans are always offsets into `full_text` as produced here, so the
map carries an extractor version stamp and a sha256 of the normalized text:
a span is only meaningful alongside those two values.
"""

from __future__ import annotations

import hashlib
import re
import unicodedata
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Any

from pypdf import PdfReader

INGEST_VERSION = "0.2.0"


def _extractor_stamp(engine: str) -> str:
    return f"target-mcp-ingest/{INGEST_VERSION} ({engine})"


EXTRACTOR_VERSION = _extractor_stamp("pypdf")  # default stamp (text input / pypdf)


class ExtractionError(ValueError):
    """A document yielded implausibly little text (e.g. a scanned or
    watermarked PDF with no usable text layer). Raised so callers fail loudly
    instead of silently assessing near-empty content."""

# Canonical sections in reading order. "other" collects everything after
# discussion (funding, COI, data availability, references-adjacent matter).
CANONICAL_SECTIONS = ("abstract", "introduction", "methods", "results", "discussion", "other")

_HEADING_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("abstract", re.compile(r"^\s*abstract\s*$", re.IGNORECASE)),
    ("introduction", re.compile(r"^\s*(?:\d+\.?\s*)?(introduction|background)\s*$", re.IGNORECASE)),
    ("methods", re.compile(r"^\s*(?:\d+\.?\s*)?(methods?|materials and methods|patients and methods|study design and methods)\s*$", re.IGNORECASE)),
    ("results", re.compile(r"^\s*(?:\d+\.?\s*)?results?\s*$", re.IGNORECASE)),
    ("discussion", re.compile(r"^\s*(?:\d+\.?\s*)?(discussion|comment)\s*$", re.IGNORECASE)),
    ("other", re.compile(r"^\s*(?:\d+\.?\s*)?(references|acknowledg(e)?ments?|funding|declarations|supplementary (material|information))\s*$", re.IGNORECASE)),
]

_PROTOCOL_TABLE_RE = re.compile(
    r"(target\s+trial\s+(specification|protocol)|specification\s+and\s+emulation"
    r"|emulation\s+of\s+the\s+target\s+trial).{0,400}?(table|tab\.)"
    r"|(table|tab\.).{0,400}?(target\s+trial\s+(specification|protocol)"
    r"|specification.{0,80}emulation)",
    re.IGNORECASE | re.DOTALL,
)
# Content/delimiter fallback: a specification↔emulation table arrives without
# the literal word "table" once ingested as text (raw-text or docx paths), but
# its column-header cells survive as delimited fields. Require a column
# separator (pipe or tab) between the paired headers so prose ("we emulated a
# target trial") does not false-positive.
_PROTOCOL_TABLE_HEADER_RE = re.compile(
    r"(target\s+trial|specification|protocol\s+component)[^\n|\t]{0,40}[|\t][^\n]{0,60}?emulat"
    r"|component[^\n|\t]{0,20}[|\t][^\n]{0,60}?(target\s+trial|emulat)",
    re.IGNORECASE,
)


def _detect_protocol_table(text: str) -> bool:
    return bool(_PROTOCOL_TABLE_RE.search(text)) or bool(_PROTOCOL_TABLE_HEADER_RE.search(text))


_FLOW_DIAGRAM_RE = re.compile(
    r"(flow\s*(diagram|chart)|study\s+flow|selection\s+of\s+(the\s+)?(study\s+)?(participants|individuals|patients)"
    r".{0,120}(figure|fig\.))|((figure|fig\.)\s*\d?.{0,120}flow)",
    re.IGNORECASE | re.DOTALL,
)


# Supplement retrieval / availability states, stamped on every SectionMap and
# carried into gating. See docs/INGESTION-AND-SCORING-DESIGN.md sec 1.4.
#   retrieved     — supplement fetched automatically and ingested
#   user_provided — supplement supplied by hand and ingested
#   none_exists   — the article is known to have no supplement
#   not_retrieved — a supplement may exist but was not obtained
#   not_checked   — single-document parse; supplement availability unknown
SUPPLEMENT_STATES = ("retrieved", "user_provided", "none_exists", "not_retrieved", "not_checked")


@dataclass
class Section:
    name: str          # canonical section
    heading: str       # heading text as matched (or "" for front matter)
    start: int         # char offset into full_text, inclusive
    end: int           # exclusive
    source: str = "main"  # "main" or "supplement:<filename>"


@dataclass
class SectionMap:
    source: str                       # path or identifier the document came from
    manuscript_id: str                # caller-supplied or derived from filename
    extractor_version: str
    text_sha256: str
    full_text: str
    sections: list[Section] = field(default_factory=list)
    n_pages: int | None = None
    protocol_table_detected: bool = False
    flow_diagram_detected: bool = False
    warnings: list[str] = field(default_factory=list)
    supplement_status: str = "not_checked"
    documents: list[dict[str, Any]] = field(default_factory=list)  # per-source metadata

    def section_text(self, name: str) -> str:
        return "\n\n".join(
            self.full_text[s.start:s.end] for s in self.sections if s.name == name
        )

    def locate(self, quote: str) -> tuple[int, int] | None:
        """Resolve a verbatim-ish quote to a (start, end) span in full_text.

        Whitespace-insensitive: both sides are collapsed before matching, then
        the match is mapped back to original offsets.
        """
        norm_text, index_map = _collapse_ws(self.full_text)
        norm_quote, _ = _collapse_ws(quote)
        if not norm_quote:
            return None
        pos = norm_text.lower().find(norm_quote.lower())
        if pos < 0:
            return None
        start = index_map[pos]
        end = index_map[pos + len(norm_quote) - 1] + 1
        return (start, end)

    def section_at(self, offset: int) -> str:
        for s in self.sections:
            if s.start <= offset < s.end:
                return s.name
        return "other"

    def source_at(self, offset: int) -> str:
        for s in self.sections:
            if s.start <= offset < s.end:
                return s.source
        return "main"

    def to_dict(self, include_text: bool = False) -> dict[str, Any]:
        d = asdict(self)
        if not include_text:
            d["full_text"] = f"<{len(self.full_text)} chars omitted>"
        return d


def _collapse_ws(text: str) -> tuple[str, list[int]]:
    """Collapse whitespace runs to single spaces; return collapsed text and a
    map from collapsed index -> original index."""
    out: list[str] = []
    index_map: list[int] = []
    in_ws = True  # swallow leading whitespace
    for i, ch in enumerate(text):
        if ch.isspace():
            if not in_ws:
                out.append(" ")
                index_map.append(i)
                in_ws = True
        else:
            out.append(ch)
            index_map.append(i)
            in_ws = False
    if out and out[-1] == " ":
        out.pop()
        index_map.pop()
    return "".join(out), index_map


def _normalize(text: str) -> str:
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # de-hyphenate line-break hyphens: "confound-\ning" -> "confounding"
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    return text


def _pypdf_pages(path: Path) -> list[str]:
    reader = PdfReader(str(path))
    return [p.extract_text() or "" for p in reader.pages]


def _pdfplumber_pages(path: Path) -> list[str] | None:
    """Second-chance extractor. Accepted-manuscript proofs and some publisher
    PDFs have a weak/missing text layer under pypdf that pdfplumber recovers.
    Returns None if pdfplumber is unavailable or fails."""
    try:
        import pdfplumber
    except ImportError:
        return None
    try:
        with pdfplumber.open(str(path)) as pdf:
            return [(pg.extract_text() or "") for pg in pdf.pages]
    except Exception:
        return None


def _too_sparse(pages: list[str]) -> bool:
    n = len(pages)
    chars = len("".join(pages).strip())
    return n >= 1 and (chars < 100 or chars < 25 * n)


def _extract_pdf(path: Path) -> tuple[list[str], str]:
    """Extract page texts via a fallback chain (pypdf -> pdfplumber). Returns
    (pages, engine)."""
    pages = _pypdf_pages(path)
    engine = "pypdf"
    if _too_sparse(pages):
        alt = _pdfplumber_pages(path)
        if alt is not None and len("".join(alt).strip()) > len("".join(pages).strip()):
            pages, engine = alt, "pdfplumber"
    return pages, engine


def _assert_plausible(text: str, n_pages: int | None, source: str) -> None:
    """Fail loudly when a paged document resolves to near-empty text, rather
    than silently proceeding to assess nothing."""
    if n_pages and len(text.strip()) < 100:
        raise ExtractionError(
            f"{source}: extracted only {len(text.strip())} characters from "
            f"{n_pages} page(s) — the PDF likely has no usable text layer "
            "(scanned or watermarked) and needs OCR. Not proceeding, to avoid "
            "assessing near-empty content."
        )


def parse_pdf(path: str | Path, manuscript_id: str | None = None) -> SectionMap:
    path = Path(path)
    pages, engine = _extract_pdf(path)
    full_text = _normalize("\n".join(pages))
    _assert_plausible(full_text, len(pages), str(path))
    return _build_map(
        full_text,
        source=str(path),
        manuscript_id=manuscript_id or path.stem,
        n_pages=len(pages),
        extractor_version=_extractor_stamp(engine),
    )


def parse_text(text: str, source: str = "<text>", manuscript_id: str = "manuscript") -> SectionMap:
    return _build_map(_normalize(text), source=source, manuscript_id=manuscript_id, n_pages=None)


def extract_file(path: str | Path) -> tuple[str, int | None]:
    """Extract raw (un-normalized) text from a supplement file by extension.
    Returns (text, n_pages). PDF via pypdf, .docx via python-docx, else UTF-8."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        pages, _engine = _extract_pdf(path)
        text = "\n".join(pages)
        _assert_plausible(text, len(pages), str(path))
        return text, len(pages)
    if suffix == ".docx":
        try:
            import docx  # python-docx
        except ImportError as e:
            raise RuntimeError("Reading .docx supplements requires python-docx") from e
        d = docx.Document(str(path))
        blocks = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                blocks.append("\t".join(c.text for c in row.cells))
        return "\n".join(blocks), None
    return path.read_text(encoding="utf-8", errors="replace"), None


_DOC_EXTENSIONS = (".pdf", ".docx", ".doc", ".txt", ".md", ".rtf")


def _looks_like_path(s: str) -> bool:
    """A single-line, reasonably short string that names a document file or
    contains a path separator is intended as a path, not as manuscript text."""
    s = s.strip()
    if not s or "\n" in s or len(s) > 400:
        return False
    return s.lower().endswith(_DOC_EXTENSIONS) or "/" in s or "\\" in s


def parse_document(path_or_text: str, manuscript_id: str | None = None) -> SectionMap:
    """Dispatch: an existing file path -> parse the file; raw text -> parse_text.

    A string that *looks like* a file path (single line, a document extension or
    a path separator) but does not resolve to an existing file raises loudly,
    rather than being silently ingested as its own ~30-character 'manuscript' —
    the most common cause is a relative path against the server's working
    directory. Pass an absolute path, or the text itself."""
    s = path_or_text
    p = Path(s)
    try:
        is_file = p.is_file()
    except OSError:  # raw text can exceed max path component length
        is_file = False
    if is_file:
        if p.suffix.lower() == ".pdf":
            return parse_pdf(p, manuscript_id)
        return parse_text(p.read_text(encoding="utf-8", errors="replace"),
                          source=str(p), manuscript_id=manuscript_id or p.stem)
    if _looks_like_path(s):
        import os
        raise FileNotFoundError(
            f"{s!r} looks like a file path but no such file exists "
            f"(server working directory: {os.getcwd()!r}). Pass an ABSOLUTE path, "
            "or pass the manuscript text itself instead of a path."
        )
    return parse_text(path_or_text, manuscript_id=manuscript_id or "manuscript")


def _build_map(full_text: str, source: str, manuscript_id: str, n_pages: int | None,
               extractor_version: str = EXTRACTOR_VERSION) -> SectionMap:
    warnings: list[str] = []
    boundaries: list[tuple[int, str, str]] = []  # (offset, canonical, heading text)

    offset = 0
    for line in full_text.split("\n"):
        stripped = line.strip()
        if stripped and len(stripped) <= 60:
            for canonical, pat in _HEADING_PATTERNS:
                if pat.match(stripped):
                    boundaries.append((offset, canonical, stripped))
                    break
        offset += len(line) + 1

    # Keep only the first occurrence of each canonical section, in document
    # order, and require order to be non-regressing (a "Results" heading
    # appearing before "Methods" is likely a running header artifact).
    seen: dict[str, int] = {}
    ordered: list[tuple[int, str, str]] = []
    rank = {name: i for i, name in enumerate(CANONICAL_SECTIONS)}
    last_rank = -1
    for off, canonical, heading in boundaries:
        if canonical in seen:
            continue
        if rank[canonical] < last_rank:
            warnings.append(f"Out-of-order heading {heading!r} at {off} ignored")
            continue
        seen[canonical] = off
        ordered.append((off, canonical, heading))
        last_rank = rank[canonical]

    sections: list[Section] = []
    if not ordered:
        warnings.append("No section headings detected; whole document mapped as 'other'")
        sections.append(Section("other", "", 0, len(full_text)))
    else:
        if ordered[0][0] > 0:
            # Front matter (title, authors, and usually the abstract when the
            # journal styles it without a literal 'Abstract' heading).
            name = "abstract" if "abstract" not in seen else "other"
            sections.append(Section(name, "", 0, ordered[0][0]))
            if name == "abstract":
                warnings.append("No 'Abstract' heading; front matter mapped as abstract")
        for i, (off, canonical, heading) in enumerate(ordered):
            end = ordered[i + 1][0] if i + 1 < len(ordered) else len(full_text)
            sections.append(Section(canonical, heading, off, end))

    missing = [s for s in ("abstract", "methods", "results") if s not in {x.name for x in sections}]
    if missing:
        warnings.append(f"Sections not detected: {missing}")

    return SectionMap(
        source=source,
        manuscript_id=manuscript_id,
        extractor_version=extractor_version,
        text_sha256=hashlib.sha256(full_text.encode("utf-8")).hexdigest(),
        full_text=full_text,
        sections=sections,
        n_pages=n_pages,
        protocol_table_detected=_detect_protocol_table(full_text),
        flow_diagram_detected=bool(_FLOW_DIAGRAM_RE.search(full_text)),
        warnings=warnings,
        supplement_status="not_checked",
        documents=[{
            "source": "main", "kind": "main", "filename": source,
            "char_start": 0, "char_end": len(full_text),
            "sha256": hashlib.sha256(full_text.encode("utf-8")).hexdigest(),
            "n_pages": n_pages,
        }],
    )


_SUPPLEMENT_SEP = "\n\n===== SUPPLEMENTARY MATERIAL: {name} =====\n\n"


def build_bundle(
    main: SectionMap,
    supplements: list[tuple[str, str, int | None]],
    supplement_status: str,
) -> SectionMap:
    """Merge a main-text SectionMap with supplement documents into one map.

    `supplements` is a list of (filename, normalized_text, n_pages). Supplement
    text is appended after the main text; each supplement becomes one Section
    (name "supplement", source "supplement:<filename>") so evidence resolved
    into it carries a truthful source locator. Structural detection (protocol
    table, flow diagram) is recomputed over the combined text, because the
    target-trial specification table commonly lives in the supplement.
    """
    if supplement_status not in SUPPLEMENT_STATES:
        raise ValueError(f"Unknown supplement_status {supplement_status!r}")
    if not supplements:
        main.supplement_status = supplement_status
        return main

    parts = [main.full_text]
    sections = [Section(s.name, s.heading, s.start, s.end, "main") for s in main.sections]
    documents = list(main.documents)
    cursor = len(main.full_text)
    for filename, text, n_pages in supplements:
        sep = _SUPPLEMENT_SEP.format(name=filename)
        text = _normalize(text)
        parts.append(sep)
        parts.append(text)
        seg_start = cursor + len(sep)
        seg_end = seg_start + len(text)
        src = f"supplement:{filename}"
        sections.append(Section("supplement", f"SUPPLEMENT: {filename}",
                                seg_start, seg_end, src))
        documents.append({
            "source": src, "kind": "supplement", "filename": filename,
            "char_start": seg_start, "char_end": seg_end,
            "sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
            "n_pages": n_pages,
        })
        cursor = seg_end

    full_text = "".join(parts)
    return SectionMap(
        source=main.source,
        manuscript_id=main.manuscript_id,
        extractor_version=EXTRACTOR_VERSION,
        text_sha256=hashlib.sha256(full_text.encode("utf-8")).hexdigest(),
        full_text=full_text,
        sections=sections,
        n_pages=main.n_pages,
        protocol_table_detected=_detect_protocol_table(full_text),
        flow_diagram_detected=bool(_FLOW_DIAGRAM_RE.search(full_text)),
        warnings=list(main.warnings),
        supplement_status=supplement_status,
        documents=documents,
    )
