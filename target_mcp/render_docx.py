"""Word (.docx) writer for the completed TARGET checklist.

Takes the dict `render.render_checklist` produces and writes a submission-ready
Word document: a title block, a completeness tally, and one table per manuscript
section — mirroring the Markdown rendering. The `enriched` view carries verdict
+ evidence columns; `official` strips to the published form's three columns
(Item / Checklist item / Location reported).

Optional logo: the published TARGET banner (bundled as assets/target-logo.png,
extracted from the official editable checklist) can head the document, but ONLY
on the `official` view — the author's self-report form. It is deliberately not
placed on the enriched view, where the official mark on a machine-generated
assessment would imply an endorsement the instrument does not carry. Off by
default; the caller opts in.

A pure serializer of an already-built report: no scoring, no spec access.
"""

from __future__ import annotations

import base64
import io
import os
from importlib import resources
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from . import render as _render
from .spec import SECTIONS

_LOGO_ASSET = "target-logo.png"


def _logo_stream() -> io.BytesIO | None:
    """The bundled TARGET banner as a fresh stream, or None if unavailable."""
    try:
        data = resources.files("target_mcp.assets").joinpath(_LOGO_ASSET).read_bytes()
    except (FileNotFoundError, ModuleNotFoundError):
        return None
    return io.BytesIO(data)

_VERDICT_RGB = {
    "reported": RGBColor(0x1B, 0x7F, 0x37),      # green
    "partial": RGBColor(0xB0, 0x6A, 0x00),       # amber
    "not_reported": RGBColor(0xA6, 0x1B, 0x1B),  # red
    "not_applicable": RGBColor(0x66, 0x66, 0x66),
    "not_assessed": RGBColor(0x88, 0x88, 0x88),
}
_ENRICHED_HEADERS = ["Item", "Checklist item", "Verdict", "Location reported",
                     "Evidence & rationale"]
_OFFICIAL_HEADERS = ["Item", "Checklist item", "Location reported"]
# Column widths in inches, per view (page is landscape; usable width ~9").
_ENRICHED_WIDTHS = [0.6, 3.3, 1.1, 1.6, 2.4]
_OFFICIAL_WIDTHS = [0.7, 6.0, 2.3]


def _set_widths(table: Any, widths_in: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_in):
            cell.width = Inches(w)


def _small(run: Any, size: int = 9) -> None:
    run.font.size = Pt(size)


def _write_item_cell(cell: Any, row: dict[str, Any], view: str) -> None:
    p = cell.paragraphs[0]
    if row["group"] and view == "enriched":
        role = "specification" if row["role"] == "specification" else "emulation"
        g = p.add_run(f"{row['group']} ")
        g.bold = True
        _small(g)
        r = p.add_run(f"({role}) — ")
        r.italic = True
        _small(r)
    body = p.add_run(row["official_text"])
    _small(body)


def _write_verdict_cell(cell: Any, row: dict[str, Any]) -> None:
    p = cell.paragraphs[0]
    run = p.add_run(row["verdict_display"])
    run.bold = True
    run.font.color.rgb = _VERDICT_RGB.get(row["verdict"], RGBColor(0, 0, 0))
    _small(run)


def _write_plain_cell(cell: Any, text: str, size: int = 9) -> None:
    run = cell.paragraphs[0].add_run(text or "")
    _small(run, size)


def write_checklist_docx(
    report: dict[str, Any], path: str, logo: bool = False,
) -> dict[str, Any]:
    """Serialize a render_checklist report to a .docx at `path`. Returns the
    absolute path, view, row count, file size, and whether the logo was applied.

    `logo=True` heads the document with the published TARGET banner, but only on
    the `official` view (the author self-report form); on the enriched view it is
    a no-op, and `logo_applied` comes back False.
    """
    view = report.get("view", "enriched")
    headers = _ENRICHED_HEADERS if view == "enriched" else _OFFICIAL_HEADERS
    widths = _ENRICHED_WIDTHS if view == "enriched" else _OFFICIAL_WIDTHS

    prov = report.get("provenance", {})
    doc = Document()
    # Self-attesting provenance in the file's core properties, so a hand-built
    # document (which cannot carry the correct stamp) is detectable at point of use.
    cp = doc.core_properties
    cp.title = _render.TITLE
    cp.author = prov.get("instrument", _render.INSTRUMENT)
    cp.category = "TARGET completeness assessment"
    cp.keywords = prov.get("stamp", "")
    cp.comments = (f"stamp {prov.get('stamp','')}; spec {prov.get('spec_version','')}; "
                   f"text_sha256 {prov.get('text_sha256','')}; "
                   f"prompt {prov.get('prompt_hash','')}")

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for side in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, side, Inches(0.6))

    logo_applied = False
    if logo and view == "official":
        stream = _logo_stream()
        if stream is not None:
            doc.add_picture(stream, width=Inches(3.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_applied = True

    doc.add_heading(_render.TITLE, level=0)
    doc.add_heading(_render.SUBTITLE, level=2)

    cite = doc.add_paragraph()
    cr = cite.add_run(report.get("attribution", {}).get("citation", _render.CITATION))
    cr.italic = True
    _small(cr, 8)

    meta = doc.add_paragraph()
    meta.add_run(f"Manuscript: {report.get('manuscript_id') or '(unnamed)'}").bold = True
    instr = doc.add_paragraph()
    _small(instr.add_run(_render.instrument_line(report)))

    comp = doc.add_paragraph()
    comp.add_run(_render.completeness_line(report["completeness"])).bold = True

    rows_by_section: dict[str, list[dict[str, Any]]] = {}
    for row in report["rows"]:
        rows_by_section.setdefault(row["section"], []).append(row)

    for sec in SECTIONS:
        section_rows = rows_by_section.get(sec)
        if not section_rows:
            continue
        doc.add_heading(_render.SECTION_TITLE[sec], level=2)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for cell, htext in zip(table.rows[0].cells, headers):
            run = cell.paragraphs[0].add_run(htext)
            run.bold = True
            _small(run)
        for row in section_rows:
            cells = table.add_row().cells
            _write_plain_cell(cells[0], row["id"])
            _write_item_cell(cells[1], row, view)
            if view == "enriched":
                _write_verdict_cell(cells[2], row)
                _write_plain_cell(cells[3], row["location"])
                _write_plain_cell(cells[4], _render._evidence_cell(row))
            else:
                _write_plain_cell(cells[2], row["location"])
        _set_widths(table, widths)

    doc.add_paragraph()
    stamp_p = doc.add_paragraph()
    sr = stamp_p.add_run(_render.provenance_line(prov))
    sr.bold = True
    _small(sr, 8)
    attrib = report.get("attribution", {})
    if attrib.get("license"):  # citation now sits under the title
        p = doc.add_paragraph()
        r = p.add_run(attrib["license"])
        r.italic = True
        _small(r, 8)

    abspath = os.path.abspath(path)
    os.makedirs(os.path.dirname(abspath) or ".", exist_ok=True)
    doc.save(abspath)
    with open(abspath, "rb") as fh:
        raw = fh.read()
    return {
        "path": abspath,
        "view": view,
        "n_rows": len(report["rows"]),
        "bytes": len(raw),
        "logo_applied": logo_applied,
        # The file lives on the server's filesystem, which a chat/agent caller
        # cannot read. Return the bytes so the caller can write the deliverable
        # on its own side; `path` is only useful for server-local/batch runs.
        "content_base64": base64.b64encode(raw).decode("ascii"),
        "content_type": ("application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document"),
        "filename": os.path.basename(abspath),
    }
