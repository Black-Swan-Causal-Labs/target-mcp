"""Minimal, dependency-light Markdown -> .docx converter for the methods doc.

Handles the constructs this document actually uses: ATX headings, pipe tables
(with an alignment separator row), blockquotes, unordered/ordered lists, thematic
breaks, and inline **bold** / *italic* / `code` / [text](url). Not a general
Markdown implementation.

Usage: python md_to_docx.py INPUT.md OUTPUT.docx
"""

from __future__ import annotations

import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_INLINE = re.compile(
    r"(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)


def add_inline(paragraph, text: str) -> None:
    """Render inline **bold**, *italic*, `code`, and [text](url) into runs."""
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            paragraph.add_run(m.group(1)).italic = True  # link text; url dropped inline
        else:
            paragraph.add_run(part)


def flush_table(doc, rows: list[list[str]], aligns: list[str]) -> None:
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci in range(ncol):
            cell = table.cell(ri, ci)
            cell.paragraphs[0].text = ""
            text = row[ci] if ci < len(row) else ""
            add_inline(cell.paragraphs[0], text)
            if ci < len(aligns):
                a = aligns[ci]
                cell.paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.RIGHT if a == "right"
                    else WD_ALIGN_PARAGRAPH.CENTER if a == "center"
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
            if ri == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()


def parse_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def is_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?\s*", line))


def aligns_from(sep: str) -> list[str]:
    out = []
    for c in parse_row(sep):
        left, right = c.startswith(":"), c.endswith(":")
        out.append("center" if left and right else "right" if right else "left")
    return out


def convert(md_path: str, docx_path: str) -> None:
    lines = open(md_path, encoding="utf-8").read().split("\n")
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    i = 0
    tbuf: list[list[str]] = []
    taligns: list[str] = []

    def flush():
        nonlocal tbuf, taligns
        flush_table(doc, tbuf, taligns)
        tbuf, taligns = [], []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # table row?
        if stripped.startswith("|"):
            if i + 1 < len(lines) and is_separator(lines[i + 1]) and not tbuf:
                tbuf.append(parse_row(stripped))
                taligns = aligns_from(lines[i + 1])
                i += 2
                continue
            if tbuf:
                tbuf.append(parse_row(stripped))
                i += 1
                continue
        elif tbuf:
            flush()

        if not stripped:
            i += 1
            continue

        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            doc.add_paragraph().add_run("_" * 40).font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            i += 1
            continue

        m = re.match(r"(#{1,6})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            p = doc.add_heading(level=min(level, 4))
            p.text = ""
            add_inline(p, m.group(2))
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, stripped.lstrip("> ").rstrip())
            i += 1
            continue

        m = re.match(r"[-*]\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, m.group(1))
            i += 1
            continue

        m = re.match(r"(\d+)\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m.group(2))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    flush()
    doc.save(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
